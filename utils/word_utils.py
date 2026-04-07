"""Funciones de apoyo para escribir y dar formato al documento Word."""

from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from config import CONFIG_COMPANIAS
from utils.text_utils import limpiar_texto_global
from utils.excel_utils import exportar_imagen_excel
import os
import re
import unicodedata

from core.extractores import (
  extraer_accidentabilidad,
  extraer_reportabilidad,
  extraer_gestion_sso,
  extraer_salud_ocupacional,
  extraer_medio_ambiente,
  extraer_asuntos_publicos,
)

# Agrega al documento el elemento indicado por su nombre.
def agregar_pie_de_pagina(doc, texto_pie):
  for section in doc.sections:
    footer = section.footer
    paragraph = footer.paragraphs[0]

    paragraph.clear()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # o CENTER si lo quieres centrado

    run = paragraph.add_run(texto_pie)
    run.font.name = "Arial"
    run.font.size = Pt(11)

# Agrega al documento el elemento indicado por su nombre.
def agregar_parrafo_sin_vineta(doc, texto, bold=False, left_indent_cm=1.27, espacio_despues=6):
  texto = limpiar_texto_global(texto)

  p = doc.add_paragraph(style="Normal AMSA")
  p.paragraph_format.line_spacing = 1.0
  p.paragraph_format.space_before = Pt(0)
  p.paragraph_format.space_after = Pt(espacio_despues)
  p.paragraph_format.left_indent = Cm(left_indent_cm)
  p.paragraph_format.first_line_indent = Cm(0)

  run = p.add_run(texto.strip())
  run.font.name = "Arial"
  run.font.size = Pt(11)
  run.bold = bold

# Agrega al documento el elemento indicado por su nombre.
def agregar_viñeta_color(doc, texto, color_punto=(0x00, 0x00, 0x00), color_texto=(0x00, 0x00, 0x00), bold=False):
    texto = limpiar_texto_global(texto)
    p = doc.add_paragraph()
    run_punto = p.add_run("● ")
    run_punto.font.color.rgb = RGBColor(*color_punto)
    run_punto.font.name = "Arial"
    run_punto.font.size = Pt(11)

    run_texto = p.add_run(texto)
    run_texto.font.color.rgb = RGBColor(*color_texto)
    run_texto.font.name = "Arial"
    run_texto.font.size = Pt(11)
    if bold:
        run_texto.bold = True

# Agrega al documento el elemento indicado por su nombre.
def agregar_viñeta_inicio_negrita(doc, texto, nivel=1, espacio_despues=6):
  if not texto:
    return
  texto = limpiar_texto_global(texto)
  nivel_norm = nivel if nivel <= 4 else 4
  p = doc.add_paragraph(style=f"Viñeta {nivel_norm}")
  p.paragraph_format.line_spacing = 1.0
  p.paragraph_format.space_before = Pt(0)
  p.paragraph_format.space_after = Pt(espacio_despues)

  partes = texto.split(":", 1)

  if len(partes) > 1:
    run_bold = p.add_run(partes[0].strip() + ": ")
    run_bold.bold = True
    run_bold.font.name = "Arial"
    run_bold.font.size = Pt(11)

    run_normal = p.add_run(partes[1].strip())
    run_normal.bold = False
    run_normal.font.name = "Arial"
    run_normal.font.size = Pt(11)
  else:
    run = p.add_run(texto.strip())
    run.bold = False
    run.font.name = "Arial"
    run.font.size = Pt(11)

# Agrega al documento el elemento indicado por su nombre.
def agregar_viñeta_con_titulo(doc, texto, nivel=1, espacio_despues=6):
  if not texto:
    return
  texto = limpiar_texto_global(texto)
  nivel_norm = nivel if nivel <= 4 else 4
  p = doc.add_paragraph(style=f"Viñeta {nivel_norm}")
  p.paragraph_format.line_spacing = 1.0
  p.paragraph_format.space_before = Pt(0)
  p.paragraph_format.space_after = Pt(espacio_despues)

  partes = texto.split(":", 1)

  if len(partes) > 1:
    titulo = partes[0].strip()
    contenido = partes[1].strip()

    run_bold = p.add_run(f"{titulo}: ")
    run_bold.bold = True
    run_bold.font.name = "Arial"
    run_bold.font.size = Pt(11)

    run_normal = p.add_run(contenido)
    run_normal.bold = False
    run_normal.font.name = "Arial"
    run_normal.font.size = Pt(11)
  else:
    run = p.add_run(texto.strip())
    run.bold = False
    run.font.name = "Arial"
    run.font.size = Pt(11)

# Agrega al documento el elemento indicado por su nombre.
def agregar_viñeta_plana(doc, texto, nivel=1, espacio_despues=6):
  if not texto:
    return
  texto = limpiar_texto_global(texto)
  nivel_norm = nivel if nivel <= 4 else 4
  p = doc.add_paragraph(style=f"Viñeta {nivel_norm}")
  p.paragraph_format.line_spacing = 1.0
  p.paragraph_format.space_before = Pt(0)
  p.paragraph_format.space_after = Pt(espacio_despues)

  run = p.add_run(texto)
  run.font.name = "Arial"
  run.font.size = Pt(11)
  run.bold = False

# Agrega al documento el elemento indicado por su nombre.
def agregar_bullet_negro_manual(doc, texto, left_indent_cm=1.27, bullet_indent_cm=0.85, espacio_despues=6, bold=False):
  texto = limpiar_texto_global(texto)
  p = doc.add_paragraph(style="Normal AMSA")
  p.paragraph_format.line_spacing = 1.0
  p.paragraph_format.space_before = Pt(0)
  p.paragraph_format.space_after = Pt(espacio_despues)
  p.paragraph_format.left_indent = Cm(left_indent_cm)
  p.paragraph_format.first_line_indent = Cm(bullet_indent_cm - left_indent_cm)

  run_bullet = p.add_run("•  ")
  run_bullet.font.name = "Arial"
  run_bullet.font.size = Pt(11)
  run_bullet.bold = False

  run_texto = p.add_run(texto)
  run_texto.font.name = "Arial"
  run_texto.font.size = Pt(11)
  run_texto.bold = bold

# Agrega al documento el elemento indicado por su nombre.
def agregar_texto_subrayado(doc, texto, left_indent_cm=0.85, espacio_despues=6, bold=True):
  texto = limpiar_texto_global(texto)
  p = doc.add_paragraph(style="Normal AMSA")
  p.paragraph_format.line_spacing = 1.0
  p.paragraph_format.space_before = Pt(0)
  p.paragraph_format.space_after = Pt(espacio_despues)
  p.paragraph_format.left_indent = Cm(left_indent_cm)
  p.paragraph_format.first_line_indent = Cm(0)

  run = p.add_run(texto)
  run.font.name = "Arial"
  run.font.size = Pt(11)
  run.bold = bold
  run.underline = True

# Agrega al documento el elemento indicado por su nombre.
def agregar_parrafo_fcab_alineado(doc, texto, bold=False, espacio_antes=False):
  texto = limpiar_texto_global(texto)

  if espacio_antes:
    doc.add_paragraph("")

  p = doc.add_paragraph()
  p.paragraph_format.line_spacing = 1.0
  p.paragraph_format.space_before = Pt(0)
  p.paragraph_format.space_after = Pt(6)
  p.paragraph_format.left_indent = Cm(0)
  p.paragraph_format.first_line_indent = Cm(0)

  run = p.add_run(texto)
  run.font.name = "Arial"
  run.font.size = Pt(11)
  run.bold = bold

# Agrega al documento el elemento indicado por su nombre.
def agregar_circulo_blanco_manual(doc, texto, left_indent_cm=1.9, bullet_indent_cm=1.45, espacio_despues=6):
  texto = limpiar_texto_global(texto).strip()

  p = doc.add_paragraph(style="Normal AMSA")
  p.paragraph_format.line_spacing = 1.0
  p.paragraph_format.space_before = Pt(0)
  p.paragraph_format.space_after = Pt(espacio_despues)
  p.paragraph_format.left_indent = Cm(left_indent_cm)
  p.paragraph_format.first_line_indent = Cm(bullet_indent_cm - left_indent_cm)

  run_bullet = p.add_run("o  ")
  run_bullet.font.name = "Arial"
  run_bullet.font.size = Pt(11)
  run_bullet.bold = False

  run_texto = p.add_run(texto)
  run_texto.font.name = "Arial"
  run_texto.font.size = Pt(11)
  run_texto.bold = False

# Agrega al documento el elemento indicado por su nombre.
def agregar_linea_acumulado(doc, texto):
  doc.add_paragraph("")
  agregar_texto(doc, texto)

# Agrega al documento el elemento indicado por su nombre.
def agregar_titulo(doc, texto, nivel=1, centrado=False, color=None):
    texto = limpiar_texto_global(texto)
    estilos_por_nivel = {1: "Título 1 AMSA", 2: "Título 2 AMSA"}
    estilo = estilos_por_nivel.get(nivel, "Título 1 AMSA")
    p = doc.add_paragraph(texto, style=estilo)
    if centrado:
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    else:
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    if color is not None:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    p.space_before = Pt(12)
    p.space_after = Pt(6)

# Agrega al documento el elemento indicado por su nombre.
def agregar_texto(doc, texto, bold=False, color=None, justificar=True):
    texto = limpiar_texto_global(texto)
    p = doc.add_paragraph(texto, style="Normal AMSA")
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    for run in p.runs:
        if bold:
            run.bold = True
        if color:
            run.font.color.rgb = RGBColor(*color)

# Agrega al documento el elemento indicado por su nombre.
def agregar_viñeta(doc, texto, nivel=1, bold=False, color=None, underline=False, espacio_despues=0):
    
    if not texto:
        return
    texto = limpiar_texto_global(texto)
    if texto.startswith("Medición calidad de aire"):
        p = doc.add_paragraph(style="Normal AMSA")
    else:
        nivel_norm = nivel if nivel <= 4 else 4
        p = doc.add_paragraph(style=f"Viñeta {nivel_norm}")

    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(espacio_despues)

    patron_especial = (
        r"(\(\w+\)|(?:El día\s)?\d{1,2}\sde\s\w+\sde\s\d{4}(?:\sa\slas\s\d{1,2}:\d{2})?(?:\.\s[A-Z]{2,4})?)"
    )

    if bold:
        run = p.add_run(texto)
        run.font.name = "Arial"
        run.font.size = Pt(11)
        run.bold = True
        if underline:
            run.underline = True
        if color:
            run.font.color.rgb = RGBColor(*color)
        return

    match_cabecera = re.search(r"^([^:]{2,40}):(\s|$)", texto)
    if match_cabecera:
        cabecera = match_cabecera.group(1) + ": "
        resto = texto[match_cabecera.end():]
        run_c = p.add_run(cabecera)
        run_c.font.name = "Arial"
        run_c.font.size = Pt(11)
        run_c.bold = True
        if underline:
            run_c.underline = True
        _escribir_texto_con_especiales(p, resto, patron_especial)
    else:
        _escribir_texto_con_especiales(p, texto, patron_especial)

# Implementa una parte específica de la lógica del informe.
def _escribir_texto_con_especiales(p, texto, patron):
    pos = 0
    for match in re.finditer(patron, texto):
        if match.start() > pos:
            run_normal = p.add_run(texto[pos:match.start()])
            run_normal.font.name = "Arial"
            run_normal.font.size = Pt(11)
            run_normal.bold = False
        texto_hallado = match.group(0)
        run_bold = p.add_run(texto_hallado)
        run_bold.font.name = "Arial"
        run_bold.font.size = Pt(11)
        run_bold.bold = True
        pos = match.end()
    if pos < len(texto):
        run_final = p.add_run(texto[pos:])
        run_final.font.name = "Arial"
        run_final.font.size = Pt(11)
        run_final.bold = False

# Agrega al documento el elemento indicado por su nombre.
def agregar_imagen(doc, ruta_imagen, ancho_cm, alto_cm=None, subtitulo=None):
    if os.path.exists(ruta_imagen):
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run()
        if alto_cm:
            run.add_picture(ruta_imagen, width=Cm(ancho_cm), height=Cm(alto_cm))
        else:
            run.add_picture(ruta_imagen, width=Cm(ancho_cm))
        if subtitulo:
            subt = doc.add_paragraph(subtitulo)
            subt.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            subt.style.font.size = Pt(10)
    else:
        errores.append(f"[ERROR] Imagen no encontrada: {ruta_imagen}")

# Agrega al documento el elemento indicado por su nombre.
def agregar_viñeta_fecha_inicial(doc, texto, nivel=1, espacio_despues=6):
  if not texto:
    return
  texto = limpiar_texto_global(texto)
  nivel_norm = nivel if nivel <= 4 else 4
  p = doc.add_paragraph(style=f"Viñeta {nivel_norm}")
  p.paragraph_format.line_spacing = 1.0
  p.paragraph_format.space_before = Pt(0)
  p.paragraph_format.space_after = Pt(espacio_despues)

  patron_fecha_inicio = re.compile(r"^\d{1,2}\sde\s\w+\sde\s\d{4}")

  match = patron_fecha_inicio.match(texto.strip())
  if match:
    fecha = match.group(0)
    resto = texto[len(fecha):]

    run_fecha = p.add_run(fecha)
    run_fecha.bold = True
    run_fecha.font.name = "Arial"
    run_fecha.font.size = Pt(11)

    run_resto = p.add_run(resto)
    run_resto.bold = False
    run_resto.font.name = "Arial"
    run_resto.font.size = Pt(11)
  else:
    run = p.add_run(texto)
    run.bold = False
    run.font.name = "Arial"
    run.font.size = Pt(11)

# Agrega al documento el elemento indicado por su nombre.
def agregar_bullet_manual(doc, texto, left_indent_cm=1.9, bullet_indent_cm=1.45, espacio_despues=6):
  texto = limpiar_texto_global(texto)
  p = doc.add_paragraph(style="Normal AMSA")
  p.paragraph_format.line_spacing = 1.0
  p.paragraph_format.space_before = Pt(0)
  p.paragraph_format.space_after = Pt(espacio_despues)
  p.paragraph_format.left_indent = Cm(left_indent_cm)
  p.paragraph_format.first_line_indent = Cm(bullet_indent_cm - left_indent_cm)

  run_bullet = p.add_run("• ")
  run_bullet.font.name = "Arial"
  run_bullet.font.size = Pt(11)
  run_bullet.bold = False

  run_texto = p.add_run(texto)
  run_texto.font.name = "Arial"
  run_texto.font.size = Pt(11)
  run_texto.bold = False


# Agrega al documento el elemento indicado por su nombre.
def agregar_viñeta_sin_negrita(doc, texto, nivel=3, espacio_despues=6):
  if not texto:
    return
  texto = limpiar_texto_global(texto)
  nivel_norm = nivel if nivel <= 4 else 4
  p = doc.add_paragraph(style=f"Viñeta {nivel_norm}")
  p.paragraph_format.line_spacing = 1.0
  p.paragraph_format.space_before = Pt(0)
  p.paragraph_format.space_after = Pt(espacio_despues)

  run = p.add_run(texto)
  run.font.name = "Arial"
  run.font.size = Pt(11)
  run.bold = False

# Agrega al documento el elemento indicado por su nombre.
def agregar_viñeta_full_bold(doc, texto, nivel=1, espacio_despues=6):
  if not texto:
    return
  texto = limpiar_texto_global(texto)
  nivel_norm = nivel if nivel <= 4 else 4
  p = doc.add_paragraph(style=f"Viñeta {nivel_norm}")
  p.paragraph_format.line_spacing = 1.0
  p.paragraph_format.space_before = Pt(0)
  p.paragraph_format.space_after = Pt(espacio_despues)

  run = p.add_run(texto)
  run.font.name = "Arial"
  run.font.size = Pt(11)
  run.bold = True

# Inserta la tabla de producción semanal de una faena como imagen.
def agregar_produccion_semana_faena(doc, clave, excel_madre):
  if not excel_madre:
    return

  cfg = CONFIG_COMPANIAS.get(clave)
  if not cfg:
    return

  agregar_titulo(doc, "Producción Semana", nivel=2)

  img_tabla = exportar_imagen_excel(
    excel_madre,
    clave,
    cfg["rango"],
    f"tabla_{clave}.png"
  )

  alto_imagen = 19.3 if clave == "CEN" else None

  agregar_imagen(doc, img_tabla, 19, alto_imagen, "")

  # 🔥 ESTE ES EL CAMBIO
  if clave == "CEN":
    doc.add_page_break()
