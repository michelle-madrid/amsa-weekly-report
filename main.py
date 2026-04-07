"""Punto de entrada para generar el informe semanal completo."""

import os

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

import state
from config import *
from utils.word_utils import *
from utils.excel_utils import *
from core.extractores import *
from core.renderers import *
from utils.excel_utils import _obtener_excel_app, _rangos_tablas_sso_backup_dinamico
from state import _workbooks_abiertos

# Genera el informe semanal completo a partir de los archivos de entrada.
def generar_informe():
    print("\n--- Generador de Informe Semanal NUEVA VERSION---")

    dia_inicio = input("Ingrese el día de inicio: ")
    mes_inicio = input("Ingrese el mes de inicio: ")
    dia_fin = input("Ingrese el día de término: ")
    mes_fin = input("Ingrese el mes de término: ")
    year = input("Ingrese el año: ")
    texto_pie = f"Semana del {dia_inicio} de {mes_inicio} al {dia_fin} de {mes_fin} {year}"

    orden_oficial = list(ORDEN_OFICIAL)
    print(f"\nOrden del informe oficial: {', '.join(orden_oficial)}")
    seleccion = input(
        "Indica las faenas a procesar separadas por coma (ej: MLP, ANT) o presiona ENTER para todas: "
    ).upper().replace(" ", "")

    if seleccion:
        faenas_activas = seleccion.split(",")
    else:
        faenas_activas = orden_oficial

    if MODO_DEBUG:
        excel_madre = RUTAS_DEBUG["excel_madre"]
        excel_indicadores = RUTAS_DEBUG["excel_indicadores"]
        carpeta_destino = RUTAS_DEBUG["carpeta_destino"]
        nombre_final=RUTAS_DEBUG["nombre_archivo"]
    else:
        excel_madre = seleccionar_archivo("Excel Base")
        excel_indicadores = seleccionar_archivo("Excel de indicadores")
        carpeta_destino = seleccionar_carpeta()
        nombre_final = input("\nEscribe el nombre del informe final: ")

    doc = Document(RUTA_PLANTILLA)
    section = doc.sections[0]
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

    p = doc.add_paragraph("Informe Semanal de Operación - Antofagasta PLC", style="Título 1 AMSA")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(24)

    for run in p.runs:
        run.font.color.rgb = RGBColor(0x12, 0x6F, 0x7A)

    resumen_texto = extraer_resumen_excel(excel_madre)
    for linea in resumen_texto.split("\n"):
        linea_limpia = linea.strip()
        if linea_limpia:
            agregar_texto(doc, linea_limpia)
            if linea_limpia.endswith("."):
                doc.add_paragraph()

    for _ in range(4):
        doc.add_paragraph()

    img_resumen = exportar_imagen_excel(excel_madre, "Grupo Minero FCAB PLAN", "A3:X34", "tabla_principal.png")
    agregar_imagen(doc, img_resumen, 19, 8.8, "")

    if INCLUIR_ESTADO_FASES_DESARROLLO:
        agregar_estado_fases_desarrollo(doc, excel_madre)

    doc.add_page_break()
    agregar_titulo(doc, "Gestión Hídrica", nivel=2)
    img_hidrica = exportar_imagen_excel(excel_madre, "Gestión Hídrica", "A3:W20", "gestion_hidrica.png")
    agregar_imagen(doc, img_hidrica, 19, 3.24, "")

    doc.add_page_break()
    agregar_titulo(doc, "Accidentabilidad", nivel=2)

    img_semanal = exportar_imagen_excel(excel_indicadores, "Informe Viernes", "A29:M41", "valor_semanal.png")
    img_mensual = exportar_imagen_excel(excel_indicadores, "Informe Viernes", "A15:M27", "valor_mensual.png")
    img_anual = exportar_imagen_excel(excel_indicadores, "Informe Viernes", "A1:M13", "valor_anual.png")

    bloques_accidentabilidad = [
        (img_semanal, "Indicadores Valor Semanal"),
        (img_mensual, "Indicadores Valor Mensual"),
        (img_anual, "Indicadores Valor Anual"),
    ]

    for img_path, texto_titulo in bloques_accidentabilidad:
        p_titulo = doc.add_paragraph()
        p_titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p_titulo.add_run(texto_titulo)
        run.bold = False
        run.font.name = "Arial"
        run.font.size = Pt(11)
        doc.add_paragraph()
        agregar_imagen(doc, img_path, 19, 4.3)
        doc.add_paragraph()

    informes = {}
    for clave in orden_oficial:
        if clave in faenas_activas:
            if MODO_DEBUG:
                ruta = RUTAS_DEBUG["informes"].get(clave)
            else:
                ruta = seleccionar_archivo(f"Informe {clave}")

            if ruta:
                informes[clave] = extraer_texto_word(ruta)
            else:
                informes[clave] = ""

    for clave in orden_oficial:
        cfg = CONFIG_COMPANIAS[clave]
        texto_compania = informes.get(clave, "")
        doc.add_page_break()
        agregar_titulo(doc, cfg["nombre"], nivel=1)
        if not texto_compania:
            agregar_texto(doc, "No solicitado.", color=(128, 128, 128))
            continue
        procesador = PROCESADORES_FAENA.get(clave)
        if procesador:
            procesador(doc, texto_compania, excel_madre)

    doc.add_page_break()
    agregar_titulo(doc, "Accidentabilidad Back-up", nivel=1)
    excel_app = _obtener_excel_app()
    wb_madre = _workbooks_abiertos.get(excel_madre)
    if wb_madre is None:
        wb_madre = excel_app.Workbooks.Open(excel_madre, UpdateLinks=0)
        _workbooks_abiertos[excel_madre] = wb_madre
    ws_sso = wb_madre.Worksheets("SSO")
    rangos_tablas_sso = _rangos_tablas_sso_backup_dinamico(ws_sso)
    for i, rango_tabla in enumerate(rangos_tablas_sso):
        nombre_img = f"accidentabilidad_{i + 1}.png"
        img_backup = exportar_imagen_excel(excel_madre, "SSO", rango_tabla, nombre_img)
        agregar_imagen(doc, img_backup, 19, None, "")
        doc.add_page_break()

    texto_pie = construir_texto_semana(dia_inicio, mes_inicio, dia_fin, mes_fin, year)
    agregar_pie_de_pagina(doc, texto_pie)
    ruta_guardado = os.path.join(carpeta_destino, f"{nombre_final}.docx")
    doc.save(ruta_guardado)
    cerrar_excels()
    print(f"Informe generado en: {ruta_guardado}")

if __name__ == "__main__":
    generar_informe()
